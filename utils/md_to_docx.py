"""
Markdown to branded Word Document Converter

Converts markdown files to Word documents using your organization template styles.
Usage: python md_to_docx.py <input.md> <output.docx>
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Path to your organization template and logo
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "quarto-brand" / "assets" / "brand-reference.docx"
LOGO_PATH = Path(__file__).parent.parent / "templates" / "quarto-brand" / "assets" / "logo.png"


def parse_markdown(md_text):
    """Parse markdown into structured elements."""
    lines = md_text.split('\n')
    elements = []
    i = 0

    # Skip YAML frontmatter
    if lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1  # Skip closing ---

    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line.strip()):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2).strip()
            elements.append({'type': 'header', 'level': level, 'text': text})
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue

        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or (lines[i].strip() and quote_lines)):
                if lines[i].strip().startswith('>'):
                    quote_lines.append(re.sub(r'^>\s*', '', lines[i]))
                elif lines[i].strip():
                    quote_lines.append(lines[i])
                else:
                    break
                i += 1
            elements.append({'type': 'blockquote', 'text': '\n'.join(quote_lines)})
            continue

        # Bullet list
        if re.match(r'^[\s]*[-*]\s+', line):
            bullets = []
            while i < len(lines) and re.match(r'^[\s]*[-*]\s+', lines[i]):
                bullet_text = re.sub(r'^[\s]*[-*]\s+', '', lines[i])
                bullets.append(bullet_text)
                i += 1
            elements.append({'type': 'bullets', 'items': bullets})
            continue

        # Numbered list
        if re.match(r'^[\s]*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                item_text = re.sub(r'^[\s]*\d+\.\s+', '', lines[i])
                items.append(item_text)
                i += 1
            elements.append({'type': 'numbered', 'items': items})
            continue

        # Checkbox list
        if re.match(r'^[\s]*-\s+\[[ x]\]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*-\s+\[[ x]\]\s+', lines[i]):
                checked = '[x]' in lines[i].lower()
                item_text = re.sub(r'^[\s]*-\s+\[[ x]\]\s+', '', lines[i])
                items.append({'text': item_text, 'checked': checked})
                i += 1
            elements.append({'type': 'checklist', 'items': items})
            continue

        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('-') and '|' not in lines[i]:
            para_lines.append(lines[i])
            i += 1
        elements.append({'type': 'paragraph', 'text': ' '.join(para_lines)})

    return elements


def clean_markdown_formatting(text):
    """Remove markdown formatting and return clean text with bold/italic markers."""
    # Handle bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Handle italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Handle links - keep text, remove URL
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Handle inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set hyperlink styling (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Standard hyperlink blue
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_formatted_text(paragraph, text):
    """Add text to paragraph with bold/italic/link formatting preserved."""
    # First, find all links and replace with placeholders
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    links = re.findall(link_pattern, text)

    # Process text with links
    parts = re.split(link_pattern, text)

    i = 0
    while i < len(parts):
        part = parts[i]

        # Check if this is link text (followed by URL)
        if i + 1 < len(parts) and i + 2 < len(parts):
            # Check if next parts look like a link pair
            potential_text = parts[i + 1] if i + 1 < len(parts) else None
            potential_url = parts[i + 2] if i + 2 < len(parts) else None

            # If current part is empty and next two are link text/url
            if part == '' and potential_text and potential_url and (potential_url.startswith('http') or potential_url.startswith('/')):
                add_hyperlink(paragraph, potential_text, potential_url)
                i += 3
                continue

        # Regular text - handle bold/italic
        if part:
            # Split by bold markers
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bpart in bold_parts:
                if bpart.startswith('**') and bpart.endswith('**'):
                    run = paragraph.add_run(bpart[2:-2])
                    run.bold = True
                elif bpart:
                    # Check for italic
                    italic_parts = re.split(r'(?<!\*)(\*[^*]+\*)(?!\*)', bpart)
                    for ipart in italic_parts:
                        if ipart and ipart.startswith('*') and ipart.endswith('*') and len(ipart) > 2:
                            run = paragraph.add_run(ipart[1:-1])
                            run.italic = True
                        elif ipart:
                            paragraph.add_run(ipart)
        i += 1


def parse_table(table_lines):
    """Parse markdown table into rows and columns."""
    rows = []
    for i, line in enumerate(table_lines):
        if i == 1 and re.match(r'^[\s|:-]+$', line):
            continue  # Skip separator line
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells from edges
        if cells:
            rows.append(cells)
    return rows


def create_branded_document(elements, output_path, title=None, footer_text=None, frontmatter=None):
    """Create Word document with your organization styling."""

    # Load template
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        # Clear template content but keep styles and sections
        # Keep sectPr (section properties) elements
        for element in doc.element.body[:]:
            if element.tag.endswith('sectPr'):
                continue  # Keep section properties
            doc.element.body.remove(element)
    else:
        doc = Document()

    section = doc.sections[0]

    # Add logo to header (smaller)
    if LOGO_PATH.exists():
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(0.9))

    # Add header text if specified in frontmatter
    if frontmatter and frontmatter.get('header_text'):
        header = section.header
        header_text_para = header.add_paragraph()
        header_text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lines = frontmatter['header_text'].strip().split('\n')
        for i, line in enumerate(lines):
            run = header_text_para.add_run(line.strip())
            run.font.size = Pt(9)
            if i < len(lines) - 1:
                header_text_para.add_run('\n')

    # Add footer
    if footer_text:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add memo header from frontmatter
    if frontmatter:
        memo_fields = []
        if frontmatter.get('to'):
            memo_fields.append(('TO:', frontmatter['to']))
        if frontmatter.get('from'):
            memo_fields.append(('FROM:', frontmatter['from']))
        if frontmatter.get('meeting_date'):
            date_str = str(frontmatter['meeting_date'])
            if frontmatter.get('meeting_time'):
                date_str += f" | {frontmatter['meeting_time']}"
            memo_fields.append(('DATE:', date_str))
        if frontmatter.get('subject'):
            memo_fields.append(('RE:', frontmatter['subject']))

        if memo_fields:
            # Create memo header as tight borderless table
            table = doc.add_table(rows=len(memo_fields), cols=2)
            table.autofit = True

            # Set tight column widths
            table.columns[0].width = Inches(0.7)
            table.columns[1].width = Inches(5.5)

            for i, (label, value) in enumerate(memo_fields):
                row = table.rows[i]
                row.height_rule = 1  # WD_ROW_HEIGHT_RULE.EXACTLY equivalent

                label_cell = row.cells[0]
                value_cell = row.cells[1]

                # Reduce cell margins
                for cell in [label_cell, value_cell]:
                    cell._tc.get_or_add_tcPr()

                label_para = label_cell.paragraphs[0]
                label_para.paragraph_format.space_after = Pt(0)
                label_para.paragraph_format.space_before = Pt(0)
                label_run = label_para.add_run(label)
                label_run.bold = True

                value_para = value_cell.paragraphs[0]
                value_para.paragraph_format.space_after = Pt(0)
                value_para.paragraph_format.space_before = Pt(0)
                value_para.add_run(value)

            # Add spacing after memo header
            doc.add_paragraph()

    # Process elements
    for elem in elements:
        if elem['type'] == 'header':
            level = elem['level']
            text = clean_markdown_formatting(elem['text'])

            if level == 1:
                # Main title - your organization Header
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['your organization Header']
                except:
                    p.style = 'Heading 1'
                add_formatted_text(p, elem['text'])
            elif level == 2:
                # Section header - your organization Headers
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['your organization Headers']
                except:
                    p.style = 'Heading 2'
                add_formatted_text(p, elem['text'])
            elif level == 3:
                # Sub-section - your organization Sub Title
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['your organization Sub Title']
                except:
                    p.style = 'Heading 3'
                add_formatted_text(p, elem['text'])
            else:
                # Lower level headers - fall back to your organization Sub Title if specific heading not available
                p = doc.add_paragraph()
                try:
                    p.style = f'Heading {min(level, 9)}'
                except KeyError:
                    try:
                        p.style = doc.styles['your organization Sub Title']
                    except:
                        p.style = 'Heading 3'
                add_formatted_text(p, elem['text'])

        elif elem['type'] == 'paragraph':
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['your organization Paragraph']
            except:
                p.style = 'Normal'
            add_formatted_text(p, elem['text'])

        elif elem['type'] == 'bullets':
            for item in elem['items']:
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['your organization Bullets']
                except:
                    p.style = 'List Bullet'
                add_formatted_text(p, item)

        elif elem['type'] == 'numbered':
            for idx, item in enumerate(elem['items'], 1):
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['List Paragraph']
                except:
                    pass
                run = p.add_run(f"{idx}. ")
                run.bold = True
                add_formatted_text(p, item)

        elif elem['type'] == 'checklist':
            for item in elem['items']:
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['your organization Bullets']
                except:
                    p.style = 'List Bullet'
                checkbox = '☑' if item['checked'] else '☐'
                p.add_run(f"{checkbox} ")
                add_formatted_text(p, item['text'])

        elif elem['type'] == 'blockquote':
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Intense Quote']
            except:
                p.style = 'Quote'
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, elem['text'])

        elif elem['type'] == 'table':
            rows = parse_table(elem['lines'])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        if j < len(table.rows[i].cells):
                            cell = table.rows[i].cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            try:
                                p.style = doc.styles['Table Text']
                            except:
                                pass
                            add_formatted_text(p, cell_text)

                            # Bold header row
                            if i == 0:
                                for run in p.runs:
                                    run.bold = True

        elif elem['type'] == 'hr':
            # Add a subtle line break / separator
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    doc.save(str(output_path))
    return output_path


def extract_frontmatter(md_text):
    """Extract YAML frontmatter as dict."""
    import yaml
    lines = md_text.split('\n')
    if lines[0].strip() == '---':
        end_idx = 1
        while end_idx < len(lines) and lines[end_idx].strip() != '---':
            end_idx += 1
        frontmatter = '\n'.join(lines[1:end_idx])
        try:
            data = yaml.safe_load(frontmatter)
            return data or {}
        except:
            pass
    return {}


def convert_md_to_docx(input_path, output_path=None):
    """Main conversion function."""
    input_path = Path(input_path)

    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)

    # Read markdown
    md_text = input_path.read_text(encoding='utf-8')

    # Extract frontmatter
    frontmatter = extract_frontmatter(md_text)

    # Build footer text - use custom footer if specified, otherwise auto-generate
    if frontmatter.get('footer'):
        footer_text = frontmatter['footer']
    else:
        meeting_date = frontmatter.get('meeting_date', '')
        doc_type = frontmatter.get('type', 'Brief').replace('-', ' ').title()
        footer_text = f"Your Organization | {doc_type} | {meeting_date}" if meeting_date else f"Your Organization | {doc_type}"

    # Parse markdown
    elements = parse_markdown(md_text)

    # Create document
    create_branded_document(elements, output_path, footer_text=footer_text, frontmatter=frontmatter)

    print(f"Created: {output_path}")
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    convert_md_to_docx(input_file, output_file)
